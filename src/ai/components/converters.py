from __future__ import annotations

import mimetypes
from io import BytesIO
from pathlib import Path
from typing import Union

import openpyxl
from docx import Document as DocxDocument
from haystack import Document, component
from haystack.dataclasses import ByteStream
from pypdf import PdfReader

Source = Union[str, Path, ByteStream]


@component
class XlsxConverter:
    """Converts XLSX files into Haystack Documents using openpyxl.

    Each sheet becomes a single Document whose content is the serialized rows.
    Metadata includes the source filename and sheet name.
    """

    def __init__(self) -> None:
        pass

    @component.output_types(documents=list[Document])
    def run(self, sources: list[Source]) -> dict[str, list[Document]]:
        documents: list[Document] = []
        for source in sources:
            if isinstance(source, ByteStream):
                data = source.data
                file_name = source.meta.get("file_name", "unknown.xlsx")
                source_meta = source.meta
            else:
                path = Path(source)
                data = path.read_bytes()
                file_name = path.name
                source_meta = {}

            workbook = openpyxl.load_workbook(
                BytesIO(data), read_only=True, data_only=True
            )
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                rows: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(
                        str(cell) if cell is not None else "" for cell in row
                    )
                    if row_text.strip():
                        rows.append(row_text)
                content = "\n".join(rows)
                doc = Document(
                    content=content,
                    meta={
                        **source_meta,
                        "file_name": file_name,
                        "sheet_name": sheet_name,
                        "content_type": "text",
                    },
                )
                documents.append(doc)
            workbook.close()
        return {"documents": documents}


class UnsupportedFileTypeError(ValueError):
    pass


@component
class SimpleFileConverter:
    @staticmethod
    def _source(source: Source) -> tuple[bytes, str, str | None, dict]:
        if isinstance(source, ByteStream):
            return (
                source.data,
                source.meta.get("file_name", "unknown"),
                source.mime_type,
                source.meta,
            )
        path = Path(source)
        return path.read_bytes(), path.name, mimetypes.guess_type(path.name)[0], {}

    @component.output_types(documents=list[Document])
    def run(self, sources: list[Source]) -> dict[str, list[Document]]:
        documents: list[Document] = []
        for source in sources:
            data, file_name, mime_type, source_meta = self._source(source)
            suffix = Path(file_name).suffix.lower()
            meta = {**source_meta, "file_name": file_name, "content_type": "text"}
            if (
                mime_type
                and mime_type.startswith("text/")
                or suffix
                in {
                    ".csv",
                    ".json",
                    ".md",
                    ".txt",
                    ".xml",
                }
            ):
                documents.append(Document(content=data.decode("utf-8"), meta=meta))
            elif mime_type == "application/pdf" or suffix == ".pdf":
                content = "\n\n".join(
                    page.extract_text() or "" for page in PdfReader(BytesIO(data)).pages
                )
                documents.append(Document(content=content, meta=meta))
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or suffix == ".docx"
            ):
                document = DocxDocument(BytesIO(data))
                lines = [paragraph.text for paragraph in document.paragraphs]
                for table in document.tables:
                    lines.extend(
                        "\t".join(cell.text for cell in row.cells) for row in table.rows
                    )
                documents.append(Document(content="\n".join(lines), meta=meta))
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                or suffix == ".xlsx"
            ):
                documents.extend(XlsxConverter().run(sources=[source])["documents"])
            else:
                raise UnsupportedFileTypeError(file_name)
        return {"documents": documents}
